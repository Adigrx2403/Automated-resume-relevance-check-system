import os
import re
import fitz  # PyMuPDF
from docx import Document
from typing import Optional, Dict, Any
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Utility class for extracting text from PDF and DOCX files."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file using PyMuPDF.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
                
            doc.close()
            return DocumentParser._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return DocumentParser._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from file based on extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text as string
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return ""
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return DocumentParser.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return DocumentParser.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return ""
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]]', ' ', text)
        
        # Remove extra spaces
        text = ' '.join(text.split())
        
        return text.strip()
    
    @staticmethod
    def extract_structured_info(text: str) -> Dict[str, Any]:
        """
        Extract structured information from text (emails, phones, etc.).
        
        Args:
            text: Extracted text
            
        Returns:
            Dictionary with structured information
        """
        info = {
            'emails': [],
            'phones': [],
            'urls': [],
            'skills_section': '',
            'education_section': '',
            'experience_section': ''
        }
        
        # Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        info['emails'] = re.findall(email_pattern, text)
        
        # Extract phone numbers
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        info['phones'] = re.findall(phone_pattern, text)
        
        # Extract URLs
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        info['urls'] = re.findall(url_pattern, text)
        
        # Extract sections (basic keyword-based approach)
        text_lower = text.lower()
        
        # Skills section
        skills_keywords = ['skills', 'technical skills', 'technologies', 'competencies']
        for keyword in skills_keywords:
            if keyword in text_lower:
                start_idx = text_lower.find(keyword)
                # Extract next 500 characters as skills section
                info['skills_section'] = text[start_idx:start_idx + 500]
                break
        
        # Education section
        education_keywords = ['education', 'academic', 'degree', 'university', 'college']
        for keyword in education_keywords:
            if keyword in text_lower:
                start_idx = text_lower.find(keyword)
                info['education_section'] = text[start_idx:start_idx + 300]
                break
        
        # Experience section
        experience_keywords = ['experience', 'work history', 'employment', 'professional']
        for keyword in experience_keywords:
            if keyword in text_lower:
                start_idx = text_lower.find(keyword)
                info['experience_section'] = text[start_idx:start_idx + 800]
                break
        
        return info
